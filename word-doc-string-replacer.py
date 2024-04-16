import os
import docx

# define a function that takes in the file path, old string, and new string
# and replaces the old string with the new string in the document
# and saves the document back to the file path
def replace_string_in_doc(filename, old_string, new_string):

    doc = docx.Document(filename)

    is_replaced = False

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text.find(old_string) != -1:
            #print("Replacing "+old_string+" in: "+filename)
            text = text.replace(old_string, new_string)
            is_replaced = True

        paragraph.text = text
    
    if is_replaced:
        print("Replaced "+old_string+" in: "+filename)

    doc.save(filename)

# define a function to remove header and footer from the word doc and save the file
def remove_header_footer(filename):
    doc = docx.Document(filename)

    # delete the header and footer
    for section in doc.sections:
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True
        section.header_distance = 0
        section.footer_distance = 0

    # save the file
    doc.save(filename)

def main(folder_path):
    for file_name in os.listdir(folder_path):
        if not (file_name.endswith(".docx") or file_name.endswith(".doc") or file_name.endswith(".DOCX") or file_name.endswith(".DOC")):
            print("Skipping file: " + file_name)
            continue

        file_path = os.path.join(folder_path, file_name)
        # call the replace_string_in_doc function with the file path, old and new strings
        try:
            #clean up header and footer
            remove_header_footer(file_path)
            replace_string_in_doc(file_path, "AMAZON", "EXAMPLE CORP")
            replace_string_in_doc(file_path, "Amazon", "Example Corp")
            replace_string_in_doc(file_path, "Amazon.com", "ExampleCorp.com")
            replace_string_in_doc(file_path, "AMAZON.com", "EXAMPLECORP.com.")
            replace_string_in_doc(file_path, "Web Services", "IT Services")
            replace_string_in_doc(file_path, "Twitch", "Example Corp Subsidiary One")
            replace_string_in_doc(file_path, "TWITCH", "EXAMPLE CORP SUBSIDIARY ONE")
            replace_string_in_doc(file_path, "Audible", "Example Subsidiary Two")
            replace_string_in_doc(file_path, "AUDIBLE", "SUBSIDIARY TWO")
        except Exception as e:
            print("Error in file: " + file_name + ": " + str(e))

# call the main function with the folder path and the old and new strings
main("/Users/gosaikat/Downloads/Offer_Letter_Templates/python-generated")